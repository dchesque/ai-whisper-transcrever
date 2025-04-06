import os
import datetime
import re
from fpdf import FPDF
from ..config.config import Config

def save_to_pdf(text, original_filename, language=None):
    """
    Salva a transcrição em um arquivo PDF
    
    Args:
        text: Texto da transcrição
        original_filename: Nome original do arquivo
        language: Idioma detectado (opcional)
        
    Returns:
        str: Caminho do arquivo PDF gerado
    """
    try:
        # Remover a extensão do arquivo original
        base_filename = os.path.splitext(os.path.basename(original_filename))[0]
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        pdf_filename = f"transcricao_{base_filename}_{timestamp}.pdf"
        pdf_path = os.path.join(Config.get_pdf_dir(), pdf_filename)
        
        # Criar o PDF
        pdf = FPDF()
        pdf.add_page()
        
        # Configurar fonte
        pdf.set_font("Arial", size=12)
        
        # Adicionar título
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt="Transcrição", ln=True, align='C')
        
        # Adicionar detalhes
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(200, 10, txt=f"Arquivo original: {original_filename}", ln=True)
        
        if language:
            pdf.cell(200, 10, txt=f"Idioma detectado: {language}", ln=True)
            
        pdf.cell(200, 10, txt=f"Data: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", ln=True)
        
        # Linha separadora
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(5)
        
        # Adicionar texto da transcrição
        pdf.set_font("Arial", size=11)
        
        # Dividir o texto em parágrafos
        paragraphs = text.split('\n\n')
        
        for paragraph in paragraphs:
            # Dividir parágrafos em linhas
            lines = paragraph.split('\n')
            
            for line in lines:
                # Calcular largura efetiva da página
                effective_width = 190
                line = line.strip()
                
                if not line:  # Se a linha estiver vazia, adicionar um espaço
                    pdf.ln(5)
                    continue
                
                # Definir largura máxima de caracteres por linha
                chars_per_line = 90  # Valor aproximado para fonte Arial 11
                
                # Dividir linhas muito longas
                while len(line) > 0:
                    # Se a linha couber completamente, adicionar e sair
                    if len(line) <= chars_per_line:
                        pdf.multi_cell(effective_width, 5, txt=line)
                        line = ""
                    else:
                        # Encontrar o último espaço antes do limite
                        cut_position = line[:chars_per_line].rfind(' ')
                        if cut_position <= 0:  # Se não encontrar espaço, cortar no limite
                            cut_position = chars_per_line
                        
                        pdf.multi_cell(effective_width, 5, txt=line[:cut_position])
                        line = line[cut_position:].lstrip()
            
            # Espaço entre parágrafos
            pdf.ln(3)
        
        # Adicionar rodapé
        pdf.set_y(-15)
        pdf.set_font('Arial', 'I', 8)
        pdf.cell(0, 10, f'Transcrever v{Config.APP_VERSION} - Página {pdf.page_no()}', 0, 0, 'C')
        
        # Salvar o PDF
        pdf.output(pdf_path)
        
        return pdf_path
    except Exception as e:
        raise Exception(f"Erro ao criar PDF: {str(e)}")

def save_to_txt(text, original_filename, language=None):
    """
    Salva a transcrição em um arquivo TXT
    
    Args:
        text: Texto da transcrição
        original_filename: Nome original do arquivo
        language: Idioma detectado (opcional)
        
    Returns:
        str: Caminho do arquivo TXT gerado
    """
    try:
        # Remover a extensão do arquivo original
        base_filename = os.path.splitext(os.path.basename(original_filename))[0]
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        txt_filename = f"transcricao_{base_filename}_{timestamp}.txt"
        txt_path = os.path.join(Config.get_pdf_dir(), txt_filename)  # usar o mesmo diretório dos PDFs
        
        with open(txt_path, 'w', encoding='utf-8') as f:
            # Adicionar cabeçalho
            f.write(f"TRANSCRIÇÃO\n")
            f.write(f"==========\n\n")
            f.write(f"Arquivo original: {original_filename}\n")
            
            if language:
                f.write(f"Idioma detectado: {language}\n")
                
            f.write(f"Data: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n")
            f.write(f"\n{'-' * 80}\n\n")
            
            # Adicionar texto da transcrição
            f.write(text)
            
            # Adicionar rodapé
            f.write(f"\n\n{'-' * 80}\n")
            f.write(f"Gerado com Transcrever v{Config.APP_VERSION}")
            
        return txt_path
    except Exception as e:
        raise Exception(f"Erro ao criar arquivo TXT: {str(e)}")

def save_to_srt(text, original_filename, language=None):
    """
    Salva a transcrição em um arquivo SRT (legendas)
    
    Args:
        text: Texto da transcrição
        original_filename: Nome original do arquivo
        language: Idioma detectado (opcional)
        
    Returns:
        str: Caminho do arquivo SRT gerado
    """
    try:
        # Remover a extensão do arquivo original
        base_filename = os.path.splitext(os.path.basename(original_filename))[0]
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        srt_filename = f"transcricao_{base_filename}_{timestamp}.srt"
        srt_path = os.path.join(Config.get_pdf_dir(), srt_filename)  # usar o mesmo diretório dos PDFs
        
        # Dividir o texto em sentenças
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        with open(srt_path, 'w', encoding='utf-8') as f:
            for i, sentence in enumerate(sentences, 1):
                if not sentence.strip():
                    continue
                    
                # Calcular tempos aproximados (3 segundos por sentença)
                start_time = (i - 1) * 3
                end_time = i * 3
                
                # Formatar tempos para o formato SRT
                start_formatted = format_srt_time(start_time)
                end_formatted = format_srt_time(end_time)
                
                # Escrever entrada SRT
                f.write(f"{i}\n")
                f.write(f"{start_formatted} --> {end_formatted}\n")
                f.write(f"{sentence.strip()}\n\n")
            
        return srt_path
    except Exception as e:
        raise Exception(f"Erro ao criar arquivo SRT: {str(e)}")

def format_srt_time(seconds):
    """Formata segundos para o formato de tempo SRT (HH:MM:SS,mmm)"""
    hours, remainder = divmod(seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    milliseconds = int((seconds - int(seconds)) * 1000)
    return f"{int(hours):02d}:{int(minutes):02d}:{int(seconds):02d},{milliseconds:03d}"

def save_to_docx(text, original_filename, language=None):
    """
    Salva a transcrição em um arquivo DOCX
    
    Args:
        text: Texto da transcrição
        original_filename: Nome original do arquivo
        language: Idioma detectado (opcional)
        
    Returns:
        str: Caminho do arquivo DOCX gerado
    """
    try:
        # Esta função requer a biblioteca python-docx
        # Se não estiver instalada, retornar erro
        try:
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("A biblioteca 'python-docx' não está instalada. Use 'pip install python-docx'.")
        
        # Remover a extensão do arquivo original
        base_filename = os.path.splitext(os.path.basename(original_filename))[0]
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_filename = f"transcricao_{base_filename}_{timestamp}.docx"
        docx_path = os.path.join(Config.get_pdf_dir(), docx_filename)  # usar o mesmo diretório dos PDFs
        
        # Criar o documento
        doc = Document()
        
        # Configurar margens da página
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        # Adicionar título
        title = doc.add_heading('Transcrição', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar detalhes
        doc.add_paragraph(f"Arquivo original: {original_filename}")
        if language:
            doc.add_paragraph(f"Idioma detectado: {language}")
        doc.add_paragraph(f"Data: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        
        # Adicionar linha separadora
        doc.add_paragraph("_" * 50)
        
        # Adicionar texto da transcrição
        for paragraph in text.split('\n\n'):
            p = doc.add_paragraph()
            p.add_run(paragraph)
        
        # Adicionar rodapé
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Gerado com Transcrever v{Config.APP_VERSION}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Salvar o documento
        doc.save(docx_path)
        
        return docx_path
    except ImportError as e:
        # Caso a biblioteca não esteja instalada, salvar como TXT como fallback
        return save_to_txt(text, original_filename, language)
    except Exception as e:
        raise Exception(f"Erro ao criar arquivo DOCX: {str(e)}") 